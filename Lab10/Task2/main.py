import docx, docx.shared

doc = docx.Document("task2.docx")

doc.add_picture('picture.jpg', width=docx.shared.Cm(10))
caption = doc.add_paragraph("Рисунок 1. ATMega8535")
caption.style = 'Caption'
doc.save('picture.docx')
