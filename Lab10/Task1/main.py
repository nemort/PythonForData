import docx

doc = docx.Document()
doc.add_paragraph("В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:")


doc.add_paragraph('Флеш-память: используется для хранения скетчей.',
                  style='List Bullet 2')

paragraph = doc.add_paragraph(style='List Bullet 2')
paragraph.add_run('ОЗУ (')
paragraph.add_run('SRAM').bold = True
paragraph.add_run(' - ')
paragraph.add_run('static random access memory').italic = True
paragraph.add_run(', статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.')
doc.add_paragraph('EEPROM (энергонезависимая память): используется для хранения постоянной информации.',
                  style='List Bullet 2')


doc.add_paragraph("Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.")


table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

main_cols = ("ATmega168", "ATmega328", "ATmega1280", "ATmega2560")
main_rows = ("Flash (1 кБ flash-памяти занят загрузчиком)", "SRAM", "EEPROM")

values = (("16 КБайт", "32 КБайт", "128 КБайт", "256 КБайт"),
          ("1 КБайт", "2 КБайт", "8 КБайт", "8 КБайт"),
          ("512 байт", "1024 байта", "4 КБайт", "4 КБайт"))

header_columns = table.rows[0].cells
header_rows = table.columns[0].cells

for i, col in enumerate(main_cols):
    header_columns[i + 1].text = col
    header_columns[i + 1].paragraphs[0].runs[0].font.bold = True

for i, value in enumerate(values):
    row_cells = table.add_row().cells
    row_cells[0].text = main_rows[i]
    row_cells[0].paragraphs[0].runs[0].font.bold = True

    for j, val in enumerate(value):
        row_cells[j+1].text = val

doc.add_paragraph(" ")
par = doc.add_paragraph()
par.add_run("Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°C. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.").italic = True

doc.save("Task1/task1.docx")
